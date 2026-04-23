

"""

from fpdf import FPDF

pdf = FPDF()
pdf.set_margins(15, 15, 15)
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()

# Calculate usable width after margins
usable_width = pdf.w - pdf.l_margin - pdf.r_margin  # 210 - 15 - 15 = 180

pdf.set_font("Helvetica", "B", 16)
pdf.cell(usable_width, 10, "TechCorp Employee Handbook", new_x="LMARGIN", new_y="NEXT", align="C")

chapters = [
    ("Chapter 1: Company Culture",
     "TechCorp values innovation, transparency, and collaboration. Every employee is encouraged to dedicate 20% of their time to experimental projects. Our annual hackathon, TechFest, takes place every September in Austin, Texas."),
    ("Chapter 2: Benefits",
     "Full-time employees receive comprehensive health insurance including dental and vision. The company matches 401(k) contributions up to 6%. Stock options vest over a 4-year period with a 1-year cliff. Unlimited PTO is available after the first year of employment."),
    ("Chapter 3: Professional Development",
     "Each employee has a $3,000 annual learning budget for conferences, courses, and certifications. Tuition reimbursement up to $10,000 per year is available for degree programs. Internal mentorship programs are available for all levels."),
    ("Chapter 4: Code of Conduct",
     "All employees must complete annual ethics training. Conflicts of interest must be disclosed to HR within 30 days. Violation of the code of conduct may result in disciplinary action up to and including termination."),
]

for heading, body in chapters:
    pdf.ln(8)
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(usable_width, 10, heading, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(usable_width, 10, body, new_x="LMARGIN", new_y="NEXT")

pdf.output("data/employee_handbook.pdf")






"""


"""

# Run once — generates a scanned-style PDF
from fpdf import FPDF
from PIL import Image
import os

# Step 1: Create a normal PDF
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", "B", 16)
pdf.cell(200, 10, "Scanned Invoice #INV-2025-001", ln=True, align="C")
pdf.set_font("Arial", "", 12)
pdf.multi_cell(0, 10, """
# Client: Acme Corporation
# Date: February 15, 2025
# Amount: $15,750.00
# Description: Consulting services for Q1 2025
# Payment Terms: Net 30
# Status: Unpaid
""")
pdf.output("data/temp_invoice.pdf")

# Step 2: Convert to image, then back to PDF (simulates scanning)
from pdf2image import convert_from_path

images = convert_from_path("data/temp_invoice.pdf")
images[0].save("data/scanned_invoice.pdf", "PDF")
os.remove("data/temp_invoice.pdf")

print("Created: data/scanned_invoice.pdf (image-based)")


"""


"""


# Run once to generate — not part of your app
# pip install python-docx

from docx import Document

doc = Document()
doc.add_heading("TechCorp Engineering Standards", level=1)

doc.add_heading("1. Code Review Policy", level=2)
doc.add_paragraph(
    "All code changes must be reviewed by at least two engineers before merging. "
    "Pull requests must include unit tests with minimum 80% code coverage. "
    "Reviews must be completed within 24 business hours of submission. "
    "The author must not merge their own pull request."
)

doc.add_heading("2. Deployment Process", level=2)
doc.add_paragraph(
    "Production deployments are permitted Monday through Thursday between 9 AM and 2 PM ET. "
    "Friday deployments require VP of Engineering approval. "
    "All deployments must go through staging environment first with a minimum 4-hour soak period. "
    "Rollback procedures must be documented before any production deployment."
)

doc.add_heading("3. Incident Severity Levels", level=2)
doc.add_paragraph(
    "SEV-1: Complete service outage affecting all customers. Response time: 15 minutes. "
    "SEV-2: Major feature degradation affecting more than 25% of customers. Response time: 30 minutes. "
    "SEV-3: Minor feature issue affecting less than 25% of customers. Response time: 4 hours. "
    "SEV-4: Cosmetic or non-urgent issues. Response time: next business day."
)

doc.add_heading("4. Tech Stack Standards", level=2)
doc.add_paragraph(
    "Backend services must use Python 3.11+ or Go 1.21+. "
    "All APIs must follow RESTful conventions with OpenAPI 3.0 documentation. "
    "PostgreSQL is the primary database. Redis is approved for caching. "
    "New technology additions require Architecture Review Board approval."
)

doc.add_heading("5. On-Call Rotation", level=2)
doc.add_paragraph(
    "Engineers participate in on-call rotations after completing 6 months of tenure. "
    "On-call shifts are one week long, Monday 9 AM to Monday 9 AM. "
    "On-call engineers receive $500 per week stipend plus $200 per incident handled. "
    "Swapping shifts requires 48 hours advance notice and manager approval."
)

doc.save("data/engineering_standards.docx")
print("Created: data/engineering_standards.docx")

"""


"""
from docx import Document

doc = Document()
doc.add_heading("TechCorp Vendor Contract Summary", level=1)

doc.add_heading("Vendor: CloudHost Inc.", level=2)
doc.add_paragraph(
    "Contract Period: January 2025 - December 2027. "
    "Annual Value: $2.4 million. "
    "Services: Cloud infrastructure hosting (AWS managed services). "
    "SLA: 99.95% uptime guarantee with $50,000 penalty per 0.01% below threshold. "
    "Termination: 90-day written notice required. Early termination fee: 6 months of remaining contract value."
)

doc.add_heading("Vendor: SecureAuth Systems", level=2)
doc.add_paragraph(
    "Contract Period: March 2025 - February 2026 (auto-renews annually). "
    "Annual Value: $180,000. "
    "Services: Identity and access management platform, SSO integration. "
    "SLA: 99.9% uptime, 4-hour response time for critical issues. "
    "Data Handling: SOC 2 Type II certified. All data encrypted at rest and in transit."
)

doc.add_heading("Vendor: DataPipe Analytics", level=2)
doc.add_paragraph(
    "Contract Period: June 2025 - May 2027. "
    "Annual Value: $350,000. "
    "Services: Data pipeline management, ETL processing, real-time analytics dashboard. "
    "SLA: 99.5% pipeline availability, data freshness within 15 minutes. "
    "Termination: 60-day notice. No early termination fee after first 12 months."
)

doc.save("data/vendor_contracts.docx")
print("Created: data/vendor_contracts.docx")


"""


"""
# pip install openpyxl

import openpyxl

wb = openpyxl.Workbook()

# Sheet 1: Revenue Metrics
ws1 = wb.active
ws1.title = "Revenue"
ws1.append(["quarter", "department", "revenue", "target", "attainment_pct"])
ws1.append(["Q1-2025", "Engineering", 4200000, 4000000, 105.0])
ws1.append(["Q1-2025", "Sales", 8500000, 9000000, 94.4])
ws1.append(["Q1-2025", "Marketing", 1200000, 1500000, 80.0])
ws1.append(["Q2-2025", "Engineering", 4800000, 4500000, 106.7])
ws1.append(["Q2-2025", "Sales", 9200000, 9000000, 102.2])
ws1.append(["Q2-2025", "Marketing", 1600000, 1500000, 106.7])

# Sheet 2: Headcount
ws2 = wb.create_sheet("Headcount")
ws2.append(["quarter", "department", "headcount", "open_roles", "attrition_rate"])
ws2.append(["Q1-2025", "Engineering", 145, 12, 8.5])
ws2.append(["Q1-2025", "Sales", 89, 5, 12.0])
ws2.append(["Q1-2025", "Marketing", 34, 3, 6.0])
ws2.append(["Q2-2025", "Engineering", 152, 8, 7.2])
ws2.append(["Q2-2025", "Sales", 91, 7, 10.5])
ws2.append(["Q2-2025", "Marketing", 36, 2, 5.8])

wb.save("data/quarterly_metrics.xlsx")
print("Created: data/quarterly_metrics.xlsx")


"""


"""
import openpyxl

wb = openpyxl.Workbook()

# Sheet 1: Servers
ws1 = wb.active
ws1.title = "Servers"
ws1.append(["server_id", "hostname", "environment", "cpu_cores", "ram_gb", "region", "status"])
ws1.append(["SRV-001", "prod-api-01", "production", 32, 128, "us-east-1", "active"])
ws1.append(["SRV-002", "prod-api-02", "production", 32, 128, "us-east-1", "active"])
ws1.append(["SRV-003", "prod-db-01", "production", 64, 256, "us-east-1", "active"])
ws1.append(["SRV-004", "staging-api-01", "staging", 16, 64, "us-west-2", "active"])
ws1.append(["SRV-005", "dev-api-01", "development", 8, 32, "us-west-2", "active"])
ws1.append(["SRV-006", "prod-api-03", "production", 32, 128, "eu-west-1", "active"])
ws1.append(["SRV-007", "prod-cache-01", "production", 16, 64, "us-east-1", "maintenance"])

# Sheet 2: Software Licenses
ws2 = wb.create_sheet("Licenses")
ws2.append(["software", "vendor", "license_type", "seats", "annual_cost", "renewal_date"])
ws2.append(["GitHub Enterprise", "GitHub", "enterprise", 200, 42000, "2025-06-01"])
ws2.append(["Jira", "Atlassian", "cloud", 300, 63000, "2025-09-15"])
ws2.append(["Slack Business+", "Salesforce", "annual", 350, 43750, "2025-04-01"])
ws2.append(["Datadog", "Datadog", "enterprise", 50, 95000, "2025-12-01"])
ws2.append(["Figma", "Figma", "organization", 40, 18000, "2025-07-15"])

wb.save("data/infrastructure_inventory.xlsx")
print("Created: data/infrastructure_inventory.xlsx")

"""